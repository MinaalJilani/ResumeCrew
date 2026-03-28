"""
Resume document generator — converts markdown resume text to a
professionally formatted DOCX matching the template format.
"""

import re
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Constants ─────────────────────────────────────────────────────────────────

SECTION_KEYWORDS = {
    "experience", "education", "skills", "projects", "certifications",
    "summary", "objective", "publications", "awards", "languages",
    "interests", "training", "achievements", "skills and awards",
    "volunteer", "professional summary", "work experience", "technical skills",
}

PREAMBLE_SIGNALS = [
    "based on your profile", "i'll generate", "i will generate",
    "please note", "here is", "here's", "below is",
    "tailored resume", "let me", "i have created",
]

TRAILING_SIGNALS = [
    "this is a basic", "remember to tailor", "would you like",
    "feel free to", "please let me know", "you can customize",
    "i hope this", "good luck", "best of luck",
]


# ── Utility helpers ───────────────────────────────────────────────────────────

def clean_md(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    text = re.sub(r"^#{1,6}\s*", "", text)
    return text.strip()


def strip_md_links(text: str) -> str:
    """[Display](url) → Display"""
    return re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)


def is_section(line: str) -> bool:
    c = clean_md(line).lower().rstrip(":")
    return c in SECTION_KEYWORDS


def looks_like_name(text: str) -> bool:
    words = text.split()
    if not (2 <= len(words) <= 5):
        return False
    if not all(w[0].isupper() for w in words if w.isalpha()):
        return False
    lower = text.lower()
    if any(w in lower for w in ["'ll", "will", " is ", " are ", " the ", "please", "note", "based", "contact"]):
        return False
    return True


def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_right_tab(paragraph, pos_twips: int = 8640):
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


# ── Preprocessing ─────────────────────────────────────────────────────────────

def preprocess_resume(text: str) -> str:
    lines = text.strip().splitlines()

    # Strip preamble
    start = 0
    for idx, line in enumerate(lines):
        clean = clean_md(line).strip()
        if not clean:
            continue
        lower = clean.lower()
        if any(sig in lower for sig in PREAMBLE_SIGNALS):
            start = idx + 1
            continue
        if is_section(line) or looks_like_name(clean):
            start = idx
            break

    lines = lines[start:]

    # Strip trailing bot remarks
    end = len(lines)
    for idx in range(len(lines) - 1, -1, -1):
        clean = clean_md(lines[idx]).strip().lower()
        if any(sig in clean for sig in TRAILING_SIGNALS):
            end = idx
        elif clean:
            break

    return "\n".join(lines[:end]).strip()


# ── Header extraction ─────────────────────────────────────────────────────────

def extract_header(lines: list) -> tuple:
    """
    Returns (name: str, contact_tokens: list[str], body_start: int).
    - Collects all lines before the first section keyword.
    - Mines name and all contact info (email, phone, LinkedIn, GitHub, URLs).
    - Deduplicates; strips markdown links and labels.
    """
    header_lines = []
    body_start = len(lines)
    for idx, line in enumerate(lines):
        if is_section(line):
            body_start = idx
            break
        header_lines.append(line)

    # Flatten header to one string with links stripped and labels removed
    flat = " ".join(strip_md_links(l) for l in header_lines)
    flat = re.sub(r"(?i)\b(email|phone|contact information|linkedin|github|portfolio|website)\s*[:·\-]?\s*", " ", flat)
    flat = re.sub(r"[|•*]", " ", flat)

    # ── Name: first line, stripping "Contact Information" suffix ──────────────
    name = ""
    for line in header_lines:
        clean = strip_md_links(clean_md(line)).strip()
        clean = re.sub(r"(?i)\s*contact information\s*:?", "", clean).strip()
        if clean and looks_like_name(clean):
            name = clean
            break
    if not name:
        for line in header_lines:
            clean = strip_md_links(clean_md(line)).strip()
            clean = re.sub(r"(?i)\s*contact information\s*:?", "", clean).strip()
            if clean and not is_section(line):
                name = clean
                break

    # ── Contact tokens via regex ──────────────────────────────────────────────
    found = []
    seen: set = set()

    def add(val: str):
        v = val.strip().strip(".,;)")
        k = v.lower()
        if v and k not in seen and len(v) > 2:
            seen.add(k)
            found.append(v)

    # Email
    for m in re.finditer(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", flat):
        add(m.group())

    # Phone (7+ digits)
    for m in re.finditer(r"\+?\d[\d\s\-().]{5,}\d", flat):
        if len(re.sub(r"\D", "", m.group())) >= 7:
            add(m.group().strip())

    # LinkedIn
    for m in re.finditer(r"linkedin\.com/in/[\w\-]+", flat, re.I):
        add(m.group())

    # GitHub
    for m in re.finditer(r"github\.com/[\w\-]+", flat, re.I):
        add(m.group())

    # Other URLs (portfolio, personal site)
    for m in re.finditer(r"https?://[^\s,]+", flat, re.I):
        val = m.group().rstrip(")")
        if "linkedin" not in val.lower() and "github" not in val.lower():
            add(val)

    # Fallback: second header line
    if not found and len(header_lines) > 1:
        for line in header_lines[1:]:
            clean = re.sub(r"^[-|•*]\s*", "", strip_md_links(clean_md(line)).strip())
            if clean:
                add(clean)
                break

    return name, found, body_start


# ── Main generator ────────────────────────────────────────────────────────────

def generate_resume_docx(resume_text: str, candidate_name: str = "", profile_links: list = None) -> bytes:
    resume_text = preprocess_resume(resume_text)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin   = Inches(0.8)
        sec.right_margin  = Inches(0.8)

    lines = resume_text.strip().splitlines()
    name, contact_tokens, body_start = extract_header(lines)
    if candidate_name:
        name = candidate_name

    # Merge saved profile links: prepend them and deduplicate
    if profile_links:
        seen = {t.lower() for t in contact_tokens}
        merged = []
        for link in profile_links:
            if link.lower() not in seen:
                merged.append(link)
                seen.add(link.lower())
        contact_tokens = merged + contact_tokens

    # ── Name ────────────────────────────────────────────────────────────────────
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(18)
        r.font.name = "Times New Roman"

    # ── Contact line ────────────────────────────────────────────────────────────
    if contact_tokens:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(8)
        r = p.add_run("  |  ".join(contact_tokens))
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"

    # ── Body ────────────────────────────────────────────────────────────────────
    i = body_start
    while i < len(lines):
        raw = lines[i].strip()
        i += 1
        if not raw:
            continue

        clean = clean_md(raw)

        # Section header
        if is_section(raw):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(clean.title())
            r.bold = True
            r.font.size = Pt(12)
            r.font.name = "Times New Roman"
            add_bottom_border(p)
            continue

        # Bullet point
        if raw.startswith(("-", "*", "•")):
            bullet = re.sub(r"^[-*•]\s*", "", clean)
            try:
                p = doc.add_paragraph(style="List Bullet")
            except Exception:
                # Fallback if "List Bullet" style is missing in default template
                p = doc.add_paragraph()
                bullet = f"• {bullet}"
            
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Inches(0.2)
            r = p.add_run(bullet)
            r.font.size = Pt(10)
            r.font.name = "Times New Roman"
            continue

        # Line with a date range → bold title + right-aligned date
        date_match = re.search(
            r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?\s*\d{4}\s*[-–]\s*(?:\d{4}|Present))",
            clean, re.I
        )
        if date_match:
            title_part = clean[:date_match.start()].strip().rstrip("|·-–").strip()
            date_part  = date_match.group(0).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(0)
            r1 = p.add_run(title_part)
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.name = "Times New Roman"
            p.add_run("\t")
            r2 = p.add_run(date_part)
            r2.font.size = Pt(10)
            r2.font.name = "Times New Roman"
            add_right_tab(p)
            continue

        # Bold line (markdown ** or ## heading)
        if raw.startswith("**") or raw.startswith("##"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(0)
            r = p.add_run(clean)
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = "Times New Roman"
            continue

        # Plain text
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(clean)
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
